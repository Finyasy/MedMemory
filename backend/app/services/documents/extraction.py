"""Document text extraction services (PDF, Images, etc.)."""

import asyncio
import io
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING

import fitz  # PyMuPDF
from PIL import Image

from app.config import settings

if TYPE_CHECKING:
    import numpy as np


@dataclass
class ExtractionResult:
    """Result of document text extraction."""

    text: str
    page_count: int
    confidence: float | None = None
    language: str | None = None
    metadata: dict | None = None
    direct_text: str | None = None
    ocr_text: str | None = None
    used_ocr: bool = False

    @property
    def is_empty(self) -> bool:
        return not self.text or len(self.text.strip()) == 0


class DocumentExtractor(ABC):
    """Base class for document text extraction."""

    @abstractmethod
    async def extract(self, file_path: str) -> ExtractionResult:
        """Extract text from a document.

        Args:
            file_path: Path to the document file

        Returns:
            ExtractionResult with extracted text and metadata
        """
        pass

    @abstractmethod
    def supports(self, mime_type: str) -> bool:
        """Check if this extractor supports the given MIME type."""
        pass


class PDFExtractor(DocumentExtractor):
    """Extract text from PDF documents using PyMuPDF.

    Supports:
    - Text-based PDFs (direct text extraction)
    - Image-based PDFs (OCR fallback)
    - Mixed PDFs (text + embedded images)
    """

    SUPPORTED_TYPES = ["application/pdf"]

    def __init__(self, ocr_fallback: bool = True):
        """Initialize PDF extractor.

        Args:
            ocr_fallback: Whether to use OCR for image-based pages
        """
        self.ocr_fallback = ocr_fallback

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES

    async def extract(self, file_path: str) -> ExtractionResult:
        """Extract text from PDF."""
        return await asyncio.to_thread(self._extract_sync, file_path)

    def _extract_sync(self, file_path: str) -> ExtractionResult:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        doc = fitz.open(file_path)

        try:
            pages_text = []
            direct_parts = []
            ocr_parts = []
            ocr_confidences = []
            used_ocr = False
            ocr_language = None
            page_count = len(doc)

            for page_num in range(page_count):
                page = doc[page_num]

                # Try direct text extraction first
                text = page.get_text("text")

                # If no text and OCR fallback enabled, try OCR
                if not text.strip() and self.ocr_fallback:
                    ocr_result = self._ocr_page_sync(page)
                    text = ocr_result.text
                    if text.strip():
                        used_ocr = True
                        ocr_confidences.append(ocr_result.confidence or 0.0)
                        ocr_language = ocr_result.language or ocr_language
                        ocr_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                elif text.strip():
                    direct_parts.append(f"--- Page {page_num + 1} ---\n{text}")

                if text.strip():
                    pages_text.append(f"--- Page {page_num + 1} ---\n{text}")

            # Extract metadata
            metadata = self._extract_metadata(doc)

            full_text = "\n\n".join(pages_text)
            direct_text = "\n\n".join(direct_parts) if direct_parts else None
            ocr_text = "\n\n".join(ocr_parts) if ocr_parts else None
            avg_confidence = (
                sum(ocr_confidences) / len(ocr_confidences) if ocr_confidences else None
            )

            return ExtractionResult(
                text=full_text,
                page_count=page_count,
                confidence=avg_confidence,
                language=ocr_language,
                metadata=metadata,
                direct_text=direct_text,
                ocr_text=ocr_text,
                used_ocr=used_ocr,
            )

        finally:
            doc.close()

    def _ocr_page_sync(self, page: fitz.Page) -> ExtractionResult:
        """OCR a PDF page by rendering it as an image.

        Args:
            page: PyMuPDF page object

        Returns:
            ExtractionResult with text from OCR
        """
        try:
            # Render page to image
            pix = page.get_pixmap(dpi=300)
            img_bytes = pix.tobytes("png")

            # Use ImageExtractor for OCR
            image_extractor = ImageExtractor()
            result = image_extractor.extract_from_bytes_sync(img_bytes, "image/png")

            return result

        except Exception:
            return ExtractionResult(
                text="",
                page_count=1,
                confidence=0.0,
                language="eng",
                used_ocr=True,
            )

    def _extract_metadata(self, doc: fitz.Document) -> dict:
        """Extract PDF metadata."""
        metadata = doc.metadata or {}

        return {
            "title": metadata.get("title", ""),
            "author": metadata.get("author", ""),
            "subject": metadata.get("subject", ""),
            "creator": metadata.get("creator", ""),
            "producer": metadata.get("producer", ""),
            "creation_date": metadata.get("creationDate", ""),
            "modification_date": metadata.get("modDate", ""),
        }

    async def extract_page(self, file_path: str, page_number: int) -> str:
        """Extract text from a specific page.

        Args:
            file_path: Path to the PDF file
            page_number: Page number (1-indexed)

        Returns:
            Text from the specified page
        """
        return await asyncio.to_thread(self._extract_page_sync, file_path, page_number)

    def _extract_page_sync(self, file_path: str, page_number: int) -> str:
        doc = fitz.open(file_path)
        try:
            if page_number < 1 or page_number > len(doc):
                raise ValueError(f"Invalid page number: {page_number}")

            page = doc[page_number - 1]
            text = page.get_text("text")

            if not text.strip() and self.ocr_fallback:
                text = self._ocr_page_sync(page).text

            return text
        finally:
            doc.close()


class ImageExtractor(DocumentExtractor):
    """Extract text from images using OCR (Tesseract).

    NOTE: This is now a FALLBACK extractor. For images, VisionExtractor
    (MedGemma direct vision) is preferred when vision_extraction_enabled=True.
    This extractor is still used for:
    - PDF OCR fallback (image-based PDF pages)
    - Manual OCR health checks
    - Fallback when vision extraction is disabled or fails

    Supports:
    - PNG, JPEG, TIFF images
    - Scanned documents
    - Medical imaging text overlays
    """

    SUPPORTED_TYPES = [
        "image/png",
        "image/jpeg",
        "image/jpg",
        "image/tiff",
    ]

    def __init__(self, language: str = "eng"):
        """Initialize image extractor.

        Args:
            language: Tesseract language code (default: English)
        """
        self.language = language
        self._tesseract_available = self._check_tesseract()

    def _check_tesseract(self) -> bool:
        """Check if Tesseract is available."""
        try:
            import pytesseract

            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES

    async def extract(self, file_path: str) -> ExtractionResult:
        """Extract text from image using OCR."""
        return await asyncio.to_thread(self._extract_sync, file_path)

    def _extract_sync(self, file_path: str) -> ExtractionResult:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Image file not found: {file_path}")

        # Read image
        with open(file_path, "rb") as f:
            img_bytes = f.read()

        # Detect MIME type from extension
        ext = path.suffix.lower()
        mime_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".tiff": "image/tiff",
            ".tif": "image/tiff",
        }
        mime_type = mime_map.get(ext, "image/png")

        return self.extract_from_bytes_sync(img_bytes, mime_type)

    async def extract_from_bytes(
        self,
        img_bytes: bytes,
        mime_type: str,
    ) -> ExtractionResult:
        """Extract text from image bytes."""
        return await asyncio.to_thread(
            self.extract_from_bytes_sync, img_bytes, mime_type
        )

    def extract_from_bytes_sync(
        self,
        img_bytes: bytes,
        _mime_type: str,
    ) -> ExtractionResult:
        # Load image with PIL
        original_image = Image.open(io.BytesIO(img_bytes))

        if not self._tesseract_available:
            # Log warning instead of silently failing
            import logging

            logger = logging.getLogger("medmemory")
            logger.warning(
                "Tesseract OCR not available. Cannot extract text from image. "
                "Install Tesseract: brew install tesseract (macOS) or "
                "apt-get install tesseract-ocr (Linux)"
            )
            return ExtractionResult(
                text="",
                page_count=1,
                confidence=0.0,
                language=self.language,
                direct_text=None,
                ocr_text="",
                used_ocr=True,
            )

        # Step 1: Try to detect orientation using OSD
        # If OSD fails or confidence is low, we'll try all orientations
        rotation_angle = self._detect_orientation(original_image)

        # Step 2: Try all 4 orientations and pick the best one
        # This is more reliable than trusting OSD alone, especially for medical forms
        best_text = ""
        best_confidence = 0.0
        best_angle = 0

        import logging

        logger = logging.getLogger("medmemory")

        # Test all 4 orientations: 0, 90, 180, 270 degrees counter-clockwise
        test_angles = [0, 90, 180, 270]

        # If OSD detected a specific angle, try that first
        if rotation_angle != 0 and rotation_angle in test_angles:
            test_angles = [rotation_angle] + [
                a for a in test_angles if a != rotation_angle
            ]
            logger.info(
                "OSD detected rotation: %d degrees. Testing orientations...",
                rotation_angle,
            )
        else:
            logger.info(
                "OSD unclear or failed. Testing all orientations (0, 90, 180, 270)..."
            )

        for angle in test_angles:
            # Rotate image if needed
            test_image = (
                self._rotate_image(original_image, angle)
                if angle != 0
                else original_image
            )

            # Preprocess the rotated image
            preprocessed = self._preprocess_image(test_image)

            # Try OCR with this orientation
            try:
                text, confidence = self._ocr_with_tesseract_sync(preprocessed)

                # Use this result if it's better than previous best
                if confidence > best_confidence:
                    best_text = text
                    best_confidence = confidence
                    best_angle = angle

                logger.debug(
                    "Orientation %d°: confidence=%.1f%%, text_length=%d",
                    angle,
                    confidence * 100,
                    len(text),
                )
            except Exception as e:
                logger.warning("OCR failed for orientation %d°: %s", angle, e)
                continue

        if best_angle != 0:
            logger.info(
                "Best orientation: %d degrees counter-clockwise (confidence=%.1f%%)",
                best_angle,
                best_confidence * 100,
            )

        return ExtractionResult(
            text=best_text,
            page_count=1,
            confidence=best_confidence,
            language=self.language,
            direct_text=None,
            ocr_text=best_text,
            used_ocr=True,
        )

    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR results.

        Enhanced preprocessing pipeline for handwritten medical documents:
        1. Upscale to ~300 DPI equivalent
        2. Convert to grayscale
        3. Deskew (rotation correction)
        4. Noise removal
        5. Adaptive binarization

        Args:
            image: PIL Image

        Returns:
            Preprocessed image
        """
        # Convert to RGB if needed
        if image.mode != "RGB":
            image = image.convert("RGB")

        # Step 1: Upscale to approximately 300 DPI for better OCR
        # Assume most scanned docs are ~72-150 DPI, target ~2000px minimum dimension
        target_min_dimension = 2000
        if min(image.size) < target_min_dimension:
            scale = target_min_dimension / min(image.size)
            new_size = (int(image.width * scale), int(image.height * scale))
            image = image.resize(new_size, Image.Resampling.LANCZOS)

        # Use OpenCV for advanced preprocessing
        if settings.ocr_preprocess_opencv:
            try:
                import cv2
                import numpy as np

                img_array = np.array(image)

                # Step 2: Convert to grayscale
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)

                # Step 3: Deskew (rotation correction)
                gray = self._deskew_image(gray)

                # Step 4: Noise removal using bilateral filter
                # Preserves edges while removing noise (good for handwriting)
                denoised = cv2.bilateralFilter(gray, 9, 75, 75)

                # Step 5: Contrast enhancement using CLAHE
                # (Contrast Limited Adaptive Histogram Equalization)
                clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                enhanced = clahe.apply(denoised)

                # Step 6: Adaptive binarization with optimized parameters
                # Use larger block size for forms/documents with varied content
                binarized = cv2.adaptiveThreshold(
                    enhanced,
                    255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY,
                    41,  # Block size (larger for forms)
                    11,  # Constant subtracted from mean
                )

                # Step 7: Morphological operations to clean up
                # Remove small noise specks
                kernel = np.ones((2, 2), np.uint8)
                cleaned = cv2.morphologyEx(binarized, cv2.MORPH_CLOSE, kernel)

                image = Image.fromarray(cleaned)
            except Exception:
                # Fallback to simple grayscale
                image = image.convert("L")
        else:
            image = image.convert("L")

        return image

    def _deskew_image(self, gray_image: "np.ndarray") -> "np.ndarray":
        """Detect and correct image rotation/skew.

        Uses Hough Line Transform to detect dominant line angles
        and rotates the image to correct skew.

        Args:
            gray_image: Grayscale numpy array

        Returns:
            Deskewed grayscale image
        """
        import cv2
        import numpy as np

        # Detect edges
        edges = cv2.Canny(gray_image, 50, 150, apertureSize=3)

        # Detect lines using Hough Transform
        lines = cv2.HoughLinesP(
            edges,
            rho=1,
            theta=np.pi / 180,
            threshold=100,
            minLineLength=100,
            maxLineGap=10,
        )

        if lines is None or len(lines) == 0:
            return gray_image

        # Calculate angles of detected lines
        angles = []
        for line in lines:
            x1, y1, x2, y2 = line[0]
            if x2 - x1 != 0:  # Avoid division by zero
                angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))
                # Only consider near-horizontal lines (likely text lines)
                if -45 < angle < 45:
                    angles.append(angle)

        if not angles:
            return gray_image

        # Get median angle (robust to outliers)
        median_angle = np.median(angles)

        # Only correct if skew is significant but not extreme
        if abs(median_angle) < 0.5 or abs(median_angle) > 15:
            return gray_image

        # Rotate image to correct skew
        (h, w) = gray_image.shape[:2]
        center = (w // 2, h // 2)
        rotation_matrix = cv2.getRotationMatrix2D(center, median_angle, 1.0)

        # Use white background for rotation
        rotated = cv2.warpAffine(
            gray_image,
            rotation_matrix,
            (w, h),
            flags=cv2.INTER_CUBIC,
            borderMode=cv2.BORDER_REPLICATE,
        )

        return rotated

    def _detect_orientation(self, image: Image.Image) -> int:
        """Detect image orientation using Tesseract OSD (Orientation and Script Detection).

        Returns the rotation angle needed (0, 90, 180, or 270 degrees counter-clockwise).
        Tesseract OSD only detects 0°, 90°, 180°, or 270° rotations.
        """
        import re

        import pytesseract

        try:
            # Use PSM 0 (OSD only) to detect orientation
            # image_to_osd returns a string with orientation info
            osd_str = pytesseract.image_to_osd(image, config="--psm 0")

            # Parse string output: "Rotate: 90" format
            # Tesseract returns clockwise rotation, we need counter-clockwise
            match = re.search(r"Rotate:\s*(\d+)", osd_str)
            if match:
                rotation = int(match.group(1))
                # Convert clockwise to counter-clockwise
                return (360 - rotation) % 360
        except Exception:
            # If OSD fails, try alternative detection methods
            pass

        # Fallback: Use aspect ratio heuristic for common case (90° rotation)
        # Medical forms are often scanned sideways
        width, height = image.size
        if width > height * 1.3:
            # Wide image - might be rotated 90° clockwise (needs 90° counter-clockwise)
            return 90
        elif height > width * 1.3:
            # Tall image - might be rotated 90° counter-clockwise (needs 90° clockwise = 270°)
            # But we'll try 90° first as it's more common
            return 0  # Assume correct orientation

        return 0  # Default: no rotation needed

    def _rotate_image(self, image: Image.Image, angle: int) -> Image.Image:
        """Rotate image by specified angle (counter-clockwise).

        Args:
            image: PIL Image
            angle: Rotation angle in degrees (0, 90, 180, or 270)

        Returns:
            Rotated image
        """
        if angle == 0:
            return image

        # Use expand=True to avoid cropping
        return image.rotate(angle, expand=True, fillcolor="white")

    def _get_ocr_config(self, use_osd: bool = False) -> str:
        """Get optimized Tesseract configuration for medical documents.

        Args:
            use_osd: Whether to use OSD (PSM 1) for orientation detection

        Returns:
            Configuration string for pytesseract
        """
        # Medical document character whitelist
        # Includes common medical symbols, units, and punctuation
        # Note: Some Tesseract versions may not support whitelist, so we wrap in try/except
        char_whitelist = (
            "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
            "abcdefghijklmnopqrstuvwxyz"
            "0123456789"
            ".,:;/-()[]{}%°"
            " \n\t"
        )

        config_parts = []

        # Page Segmentation Mode
        if use_osd:
            # PSM 1: Automatic page segmentation with OSD
            config_parts.append("--psm 1")
        else:
            # PSM 6: Assume a single uniform block of text (good for forms)
            # This works well for structured medical documents
            config_parts.append("--psm 6")

        # Character whitelist to reduce garbage characters
        # This helps filter out noise and improves accuracy for medical documents
        # Tesseract will ignore this if not supported (some versions don't support it)
        config_parts.append(f"-c tessedit_char_whitelist={char_whitelist}")

        # Explicit DPI setting (300 DPI is optimal for OCR)
        config_parts.append("-c user_defined_dpi=300")

        # Additional optimizations for medical documents
        config_parts.append("-c preserve_interword_spaces=1")  # Keep spacing

        return " ".join(config_parts)

    def _ocr_with_tesseract_sync(
        self,
        image: Image.Image,
    ) -> tuple[str, float]:
        """Perform OCR using Tesseract with optimized configuration for medical documents.

        Note: Orientation detection and rotation should be done BEFORE calling this method
        (in extract_from_bytes_sync) to avoid double-processing.

        This method:
        1. Applies optimized Tesseract configuration for medical documents
        2. Performs OCR with confidence tracking
        3. Falls back to OSD mode if confidence is very low
        """
        import pytesseract

        # Get optimized OCR configuration
        # Try PSM 6 first (single block) for better accuracy on forms
        config = self._get_ocr_config(use_osd=False)

        # Step 3: Perform OCR with medical document optimizations
        try:
            data = pytesseract.image_to_data(
                image,
                lang=self.language,
                config=config,
                output_type=pytesseract.Output.DICT,
            )
        except Exception:
            # Fallback to default config if custom config fails
            data = pytesseract.image_to_data(
                image,
                lang=self.language,
                output_type=pytesseract.Output.DICT,
            )

        # Extract text and confidence scores
        text_parts = []
        confidences = []

        for i, conf in enumerate(data["conf"]):
            if conf > 0:  # Valid detection
                text = data["text"][i].strip()
                if text:
                    text_parts.append(text)
                    confidences.append(conf)

        full_text = " ".join(text_parts)
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

        # If confidence is very low, try OSD mode (PSM 1) as fallback
        # OSD mode includes orientation detection and may help with difficult images
        if avg_confidence < 30:
            try:
                # Retry with OSD mode (PSM 1) which includes orientation detection
                config_osd = self._get_ocr_config(use_osd=True)
                data_osd = pytesseract.image_to_data(
                    image,
                    lang=self.language,
                    config=config_osd,
                    output_type=pytesseract.Output.DICT,
                )

                text_parts_osd = []
                confidences_osd = []
                for i, conf in enumerate(data_osd["conf"]):
                    if conf > 0:
                        text = data_osd["text"][i].strip()
                        if text:
                            text_parts_osd.append(text)
                            confidences_osd.append(conf)

                if confidences_osd:
                    avg_conf_osd = sum(confidences_osd) / len(confidences_osd)
                    # Use OSD result if it's better
                    if avg_conf_osd > avg_confidence:
                        full_text = " ".join(text_parts_osd)
                        avg_confidence = avg_conf_osd
            except Exception:
                pass  # Keep original result

        return full_text, avg_confidence / 100.0  # Normalize to 0-1


class VisionExtractor(DocumentExtractor):
    """Extract structured data directly from images using MedGemma vision.

    This extractor bypasses OCR entirely and uses MedGemma's vision capabilities
    to directly read and extract structured medical information from images.
    Better accuracy for both printed and handwritten text in medical documents.
    """

    SUPPORTED_TYPES = [
        "image/png",
        "image/jpeg",
        "image/jpg",
        "image/tiff",
    ]

    def __init__(self):
        """Initialize vision extractor."""
        self.logger = logging.getLogger("medmemory")
        self._llm_service = None

    def _get_llm_service(self):
        """Lazy-load LLM service."""
        if self._llm_service is None:
            from app.services.llm import LLMService

            self._llm_service = LLMService.get_instance()
        return self._llm_service

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES

    async def extract(self, file_path: str) -> ExtractionResult:
        """Extract structured medical data directly from image using MedGemma vision."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Image file not found: {file_path}")

        # Read image bytes
        with open(file_path, "rb") as f:
            image_bytes = f.read()

        return await self.extract_from_bytes(image_bytes)

    async def extract_from_bytes(self, image_bytes: bytes) -> ExtractionResult:
        """Extract structured medical data from image bytes using MedGemma vision."""
        llm_service = self._get_llm_service()

        system_prompt = (
            "You are a medical document extraction expert. "
            "Extract all medical information from the provided document image.\n\n"
            "CRITICAL RULES:\n"
            "- ONLY extract text, numbers, and values that are explicitly visible in the image.\n"
            "- NEVER invent measurements, dimensions, or values that are not written in the document.\n"
            "- For medical images (X-rays, scans), describe what you see but do NOT invent numerical measurements unless there is a visible scale or ruler.\n"
            "- If text is unclear or unreadable, transcribe what you can see and note uncertainty.\n"
            "- Preserve all numerical values, units, and medical terminology exactly as written.\n\n"
            "Extract and include:\n"
            "- Patient information (name, DOB, ID numbers) if visible\n"
            "- Vital signs (blood pressure, heart rate, temperature, etc.) if written\n"
            "- Diagnoses and conditions if stated\n"
            "- Medications with dosages if listed\n"
            "- Lab values with units if shown\n"
            "- Clinical notes and observations if present\n"
            "- Dates and timestamps if visible\n\n"
            "Format as clear, searchable text. If something is not visible, do not invent it."
        )

        # User prompt
        user_prompt = (
            "Extract all medical information from this document image. "
            "Return structured, searchable text with all relevant medical data."
        )

        try:
            self.logger.info("Using MedGemma vision extraction for image document")
            response = await llm_service.generate_with_image(
                prompt=user_prompt,
                image_bytes=image_bytes,
                system_prompt=system_prompt,
                max_new_tokens=settings.vision_extraction_max_new_tokens,
            )

            extracted_text = response.text.strip()

            if not extracted_text:
                self.logger.warning("MedGemma vision extraction returned empty text")
                extracted_text = "[No text extracted from image]"

            self.logger.info(
                "Vision extraction completed: %d characters extracted",
                len(extracted_text),
            )

            return ExtractionResult(
                text=extracted_text,
                page_count=1,
                confidence=0.95,  # High confidence from vision model
                language="en",
                direct_text=extracted_text,
                ocr_text=None,  # Not OCR, direct vision
                used_ocr=False,
            )

        except Exception as e:
            self.logger.error(
                "MedGemma vision extraction failed: %s",
                e,
                exc_info=True,
            )
            # Return empty result - caller can fall back to OCR if needed
            return ExtractionResult(
                text="",
                page_count=1,
                confidence=0.0,
                language="en",
                direct_text=None,
                ocr_text=None,
                used_ocr=False,
            )


class DocxExtractor(DocumentExtractor):
    """Extract text from Word documents."""

    SUPPORTED_TYPES = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES

    async def extract(self, file_path: str) -> ExtractionResult:
        """Extract text from DOCX file."""
        return await asyncio.to_thread(self._extract_sync, file_path)

    def _extract_sync(self, file_path: str) -> ExtractionResult:
        from docx import Document as DocxDocument

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = DocxDocument(file_path)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return ExtractionResult(
            text="\n\n".join(paragraphs),
            page_count=1,  # DOCX doesn't have fixed pages
            direct_text="\n\n".join(paragraphs),
            ocr_text=None,
            used_ocr=False,
        )


def get_extractor(
    mime_type: str, prefer_vision: bool = None
) -> DocumentExtractor | None:
    """Get the appropriate extractor for a MIME type.

    Args:
        mime_type: MIME type of the document
        prefer_vision: Whether to prefer vision extraction for images (defaults to settings)

    Returns:
        Appropriate DocumentExtractor or None if unsupported
    """
    if prefer_vision is None:
        prefer_vision = settings.vision_extraction_enabled

    # For images, prefer vision extraction if enabled
    if prefer_vision and mime_type in VisionExtractor.SUPPORTED_TYPES:
        vision_extractor = VisionExtractor()
        if vision_extractor.supports(mime_type):
            return vision_extractor

    # Fallback to standard extractors
    # Standard extractors (ImageExtractor is fallback for images when vision disabled/fails)
    extractors = [
        PDFExtractor(),
        ImageExtractor(),  # OCR fallback for images (used when vision extraction disabled/fails)
        DocxExtractor(),
    ]

    for extractor in extractors:
        if extractor.supports(mime_type):
            return extractor

    return None
